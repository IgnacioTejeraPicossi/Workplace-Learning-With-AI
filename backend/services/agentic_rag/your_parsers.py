# your_parsers.py
from dataclasses import dataclass
from typing import List, Optional
import os, io, re

@dataclass
class Section:
    title: str
    text: str
    paragraphs: List["Paragraph"]

@dataclass
class Paragraph:
    title: Optional[str]
    text: str

# --- Utilidades ---
def _merge_small_paragraphs(paras: List[str], max_chars=1800) -> List[str]:
    out, buf = [], ""
    for p in paras:
        p = p.strip()
        if not p: continue
        if len(buf) + len(p) < max_chars:
            buf = f"{buf}\n\n{p}" if buf else p
        else:
            out.append(buf); buf = p
    if buf: out.append(buf)
    return out

def _from_plain_text(text: str, default_title="Document") -> List[Section]:
    # Detecta encabezados (#, ##) o patrones 1., 1.1 etc.
    lines = text.splitlines()
    sections: List[Section] = []
    cur_title, cur_buf = default_title, []
    header_re = re.compile(r"^(#{1,6}\s+.+|(\d+(\.\d+)*)\s+.+)$")

    for ln in lines:
        if header_re.match(ln.strip()):
            # flush
            if cur_buf:
                paras = _merge_small_paragraphs(re.split(r"\n{2,}", "\n".join(cur_buf)))
                sections.append(Section(title=cur_title, text="\n".join(cur_buf),
                                        paragraphs=[Paragraph(cur_title, p) for p in paras]))
                cur_buf = []
            cur_title = re.sub(r"^#{1,6}\s+", "", ln.strip())
        else:
            cur_buf.append(ln)
    if cur_buf:
        paras = _merge_small_paragraphs(re.split(r"\n{2,}", "\n".join(cur_buf)))
        sections.append(Section(title=cur_title, text="\n".join(cur_buf),
                                paragraphs=[Paragraph(cur_title, p) for p in paras]))
    return sections

# --- DOCX ---
def _from_docx_bytes(data: bytes) -> str:
    try:
        import docx  # python-docx
    except Exception:
        raise RuntimeError("Falta dependencia: python-docx")
    f = io.BytesIO(data)
    d = docx.Document(f)
    txt = []
    for p in d.paragraphs:
        txt.append(p.text)
    return "\n".join(txt)

# --- PDF ---
def _from_pdf_bytes(data: bytes) -> str:
    try:
        import pypdf  # PyPDF2 moderno
    except Exception:
        raise RuntimeError("Falta dependencia: pypdf")
    reader = pypdf.PdfReader(io.BytesIO(data))
    text = []
    for page in reader.pages:
        try:
            text.append(page.extract_text() or "")
        except Exception:
            text.append("")
    return "\n".join(text)

def _from_md_bytes(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")

def _from_txt_bytes(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")

def _merge_block(block: List[Section]) -> Section:
    title = block[0].title
    text = "\n\n".join(s.text for s in block)
    paras = []
    for s in block:
        paras.extend(s.paragraphs)
    return Section(title=title, text=text, paragraphs=paras)

# API principal
def parse_to_sections_and_paragraphs(content: bytes, filename: str) -> List[Section]:
    ext = os.path.splitext(filename)[1].lower()
    if ext in [".docx"]:
        text = _from_docx_bytes(content)
    elif ext in [".pdf"]:
        text = _from_pdf_bytes(content)
    elif ext in [".md"]:
        text = _from_md_bytes(content)
    elif ext in [".txt"]:
        text = _from_txt_bytes(content)
    else:
        # fallback: intenta como texto
        text = content.decode("utf-8", errors="ignore")
    
    sections = _from_plain_text(text, default_title=filename)
    
    # Recorte a ~20 mega-secciones como sugiere Agentic RAG (zero-ingestion)
    if len(sections) > 30:
        # agrupa por bloques
        merged = []
        block = []
        for s in sections:
            block.append(s)
            if len(block) >= max(1, len(sections)//20):
                merged.append(_merge_block(block))
                block = []
        if block: merged.append(_merge_block(block))
        sections = merged
    
    return sections
