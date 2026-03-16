from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Request, Query, Body
from typing import List, Tuple, Dict, Any, Optional
from io import BytesIO
import json
import re
from datetime import datetime
from fastapi.responses import StreamingResponse

try:
    from backend.prompt_loader import load_prompt_template, get_active_version
except Exception:
    # Fallback if prompt_loader is not available
    load_prompt_template = None
    get_active_version = None

try:
    from docx import Document  # python-docx
    from docx.shared import Inches
    from docx.document import Document as _Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except Exception:
    Document = None
    Inches = None
    CT_P = None
    CT_Tbl = None
    Table = None
    Paragraph = None

try:
    # Prefer the unified LLM used across the app
    from backend.llm import ask_ai_unified_sync
except Exception:
    ask_ai_unified_sync = None

router = APIRouter(prefix="/api/j-messages", tags=["J-messages Analyzer"])

try:
    # Lazy import DB to avoid import cycles during tool runs
    from backend.db import database
    j_messages_collection = database.get_collection("j_messages")
except Exception:
    j_messages_collection = None

def _normalize_category(category, categories=None) -> list:
    """
    Normalize category field to a valid list.
    Handles: list, string, legacy 'categories' field, None, and invalid values.
    Always returns a non-empty list with valid values only.
    """
    valid = {"Annet", "Bunnfisk", "Pelagisk fisk"}
    if isinstance(category, list):
        filtered = [c for c in category if c in valid]
        return filtered if filtered else ["Annet"]
    if category:
        return [category] if category in valid else ["Annet"]
    # Fallback to legacy 'categories' field
    if isinstance(categories, list):
        filtered = [c for c in categories if c in valid]
        return filtered if filtered else ["Annet"]
    if categories:
        return [categories] if categories in valid else ["Annet"]
    return ["Annet"]


def _simple_slug(value: str) -> str:
    """
    Simple slug generator to avoid external deps.
    Lowercase, replace non-alphanum with hyphens, collapse repeats.
    """
    value = value.lower()
    value = re.sub(r"\s+", "-", value)
    value = re.sub(r"[^\w\-]", "-", value)
    value = re.sub(r"-{2,}", "-", value).strip("-")
    return value or "section"


def _parse_json_relaxed(content: str) -> Dict[str, Any]:
    """
    Parse JSON; if it fails with e.g. 'Extra data' when the root is missing (content
    starts with \"data\": {...}), try wrapping in {...} or extracting the first {...} as data.
    """
    # Strip BOM and surrounding whitespace
    if content.startswith("\ufeff"):
        content = content[1:]
    content = content.strip()
    try:
        return json.loads(content)
    except json.JSONDecodeError as e1:
        if not content or content.startswith("{") or content.startswith("["):
            raise e1
        # Try wrapping as { content } (for "data": { ... } missing the outer braces)
        try:
            return json.loads("{" + content + "}")
        except json.JSONDecodeError:
            pass
        # Fallback: take from first { to last } and treat as inner object under "data"
        i = content.find("{")
        j = content.rfind("}")
        if i >= 0 and j > i:
            try:
                inner = json.loads(content[i : j + 1])
                return {"data": inner}
            except json.JSONDecodeError:
                pass
        raise e1


def _extract_text_from_json(data: Dict[str, Any]) -> str:
    """
    Get document text from JSON: raw_text, body, text; or Enonic-style data.blocks.text.text.
    HTML in blocks.text.text is stripped to plain text.
    """
    def _to_plain(h: str) -> str:
        if not h or not isinstance(h, str):
            return ""
        return re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", h)).strip()

    for key in ("raw_text", "body", "text"):
        v = data.get(key)
        if isinstance(v, str) and v.strip():
            return v.strip()

    inner = data.get("data")
    if isinstance(inner, dict):
        for key in ("raw_text", "body", "text"):
            v = inner.get(key)
            if isinstance(v, str) and v.strip():
                return v.strip()
        blocks = inner.get("blocks") or {}
        if isinstance(blocks, dict):
            text_block = blocks.get("text")
            if isinstance(text_block, dict):
                t = text_block.get("text")
                if isinstance(t, str) and t.strip():
                    return _to_plain(t)
            elif isinstance(text_block, str) and text_block.strip():
                return _to_plain(text_block)
    return ""


def read_docx_paragraphs(file_bytes: bytes) -> List[str]:
    """
    Extract paragraphs and tables from DOCX file.
    Returns a list of strings where tables are represented as special markers.
    """
    if Document is None or CT_P is None or CT_Tbl is None:
        raise HTTPException(status_code=500, detail="python-docx not available on server")
    doc = Document(BytesIO(file_bytes))
    paragraphs: List[str] = []
    
    # Get all elements in document order (paragraphs and tables)
    body = doc.element.body
    for element in body.iterchildren():
        if isinstance(element, CT_P):
            # It's a paragraph
            para = Paragraph(element, doc)
            text = (para.text or "").strip()
            if text:
                paragraphs.append(text)
        elif isinstance(element, CT_Tbl):
            # It's a table - extract it
            table = Table(element, doc)
            table_html = extract_table_as_html(table)
            if table_html:
                paragraphs.append(f"<TABLE_MARKER>{table_html}</TABLE_MARKER>")
    
    return paragraphs


def extract_table_as_html(table) -> str:
    """
    Convert a python-docx Table object to HTML table.
    
    Args:
        table: python-docx Table object
    
    Returns:
        HTML string representing the table
    """
    try:
        import html
        html_parts = ['<table style="border-collapse: collapse; width: 100%; margin: 12px 0;">']
        
        for i, row in enumerate(table.rows):
            html_parts.append('<tr>')
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                # Escape HTML special characters
                cell_text_escaped = html.escape(cell_text)
                # Use <th> for header row (first row), <td> for data rows
                tag = 'th' if i == 0 else 'td'
                style = 'border: 1px solid #ddd; padding: 8px; text-align: left;' + (' font-weight: bold; background-color: #f5f5f5;' if i == 0 else '')
                html_parts.append(f'<{tag} style="{style}">{cell_text_escaped}</{tag}>')
            html_parts.append('</tr>')
        
        html_parts.append('</table>')
        return ''.join(html_parts)
    except Exception as e:
        print(f"[J-MESSAGES] ⚠️ Failed to extract table as HTML: {e}")
        return ""


def html_table_to_docx(html_table: str, doc) -> None:
    """
    Convert an HTML table string to a python-docx Table object.
    
    Args:
        html_table: HTML string containing a <table> element
        doc: python-docx Document object to add the table to
    """
    try:
        from html.parser import HTMLParser
        import html
        
        # Extract table rows
        row_pattern = r'<tr[^>]*>(.*?)</tr>'
        rows = re.findall(row_pattern, html_table, re.DOTALL | re.IGNORECASE)
        
        if not rows:
            return
        
        # Determine number of columns from first row
        first_row = rows[0]
        cell_pattern = r'<(?:th|td)[^>]*>(.*?)</(?:th|td)>'
        first_row_cells = re.findall(cell_pattern, first_row, re.DOTALL | re.IGNORECASE)
        num_cols = len(first_row_cells)
        
        if num_cols == 0:
            return
        
        # Create DOCX table
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Light Grid Accent 1'  # Use a built-in table style
        
        # Populate table
        for row_idx, row_html in enumerate(rows):
            cells = re.findall(cell_pattern, row_html, re.DOTALL | re.IGNORECASE)
            docx_row = table.rows[row_idx]
            
            for col_idx, cell_html in enumerate(cells[:num_cols]):
                # Unescape HTML entities and strip HTML tags
                cell_text = html.unescape(cell_html)
                cell_text = re.sub(r'<[^>]+>', '', cell_text)
                cell_text = re.sub(r'\s+', ' ', cell_text).strip()
                
                docx_row.cells[col_idx].text = cell_text
                
                # Make header row bold
                if row_idx == 0:
                    for paragraph in docx_row.cells[col_idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    except Exception as e:
        print(f"[J-MESSAGES EXPORT] ⚠️ Failed to convert HTML table to DOCX: {e}")
        raise


def read_pdf_text(file_bytes: bytes) -> str:
    """
    Extract text from PDF using pypdf/PyPDF2 if available, otherwise pdfminer.six.
    """
    # 1) Try pypdf / PyPDF2
    try:
        try:
            from pypdf import PdfReader as _Reader  # type: ignore
        except Exception:
            from PyPDF2 import PdfReader as _Reader  # type: ignore
        reader = _Reader(BytesIO(file_bytes))
        parts: List[str] = []
        for page in getattr(reader, "pages", []):
            try:
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(text)
            except Exception:
                continue
        text = "\n".join(parts).strip()
        if text:
            return text
    except Exception:
        pass

    # 2) Fallback: pdfminer.six
    try:
        from pdfminer.high_level import extract_text  # type: ignore
        text = extract_text(BytesIO(file_bytes)) or ""
        return text.strip()
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"PDF processing not available. Please install one of: pypdf, PyPDF2, or pdfminer.six. Error: {e}"
        )


def split_header_body(paragraphs: List[str]) -> Tuple[str, str]:
    """
    Split at the common marker 'Forskriften lyder etter dette'.
    If not found, treat everything as body.
    """
    marker = "Forskriften lyder etter dette"
    header_lines: List[str] = []
    body_lines: List[str] = []
    found = False
    for line in paragraphs:
        if not found:
            header_lines.append(line)
            if marker in line:
                found = True
        else:
            body_lines.append(line)
    if not found:
        return "", "\n".join(paragraphs)
    return "\n".join(header_lines), "\n".join(body_lines)


def build_toc_and_body_html(body_text: str) -> Tuple[List[Dict[str, Any]], str]:
    """
    Very lightweight heading detection for first implementation:
    - Lines starting with 'Kapittel ' → H1
    - Lines starting with '§' → H2
    - Lines with <TABLE_MARKER>...</TABLE_MARKER> → extract and insert HTML table
    Everything else → <p>.
    """
    lines = [l for l in body_text.split("\n") if l.strip()]
    toc: List[Dict[str, Any]] = []
    html_parts: List[str] = []
    last_h1: Dict[str, Any] = {}
    used_anchors: Dict[str, int] = {}

    def ensure_unique(anchor: str) -> str:
        """
        Guarantee unique anchors to avoid duplicate ids and React key warnings.
        If an anchor repeats, suffix with -2, -3, ...
        """
        count = used_anchors.get(anchor, 0)
        if count == 0:
            used_anchors[anchor] = 1
            return anchor
        count += 1
        used_anchors[anchor] = count
        return f"{anchor}-{count}"

    for raw in lines:
        line = raw.strip()
        # Check if this line contains a table marker
        if "<TABLE_MARKER>" in line and "</TABLE_MARKER>" in line:
            # Extract table HTML from marker
            table_match = re.search(r'<TABLE_MARKER>(.*?)</TABLE_MARKER>', line, re.DOTALL)
            if table_match:
                table_html = table_match.group(1).strip()
                html_parts.append(table_html)
            else:
                # Fallback: try to extract any HTML table from the line
                table_match = re.search(r'<table.*?</table>', line, re.DOTALL | re.IGNORECASE)
                if table_match:
                    html_parts.append(table_match.group(0))
                else:
                    # If no table found, just add as paragraph
                    html_parts.append(f"<p>{line}</p>")
        elif line.startswith("Kapittel "):
            anchor = ensure_unique(_simple_slug(line))
            item = {"level": 1, "title": line, "anchor": anchor, "children": []}
            toc.append(item)
            last_h1 = item
            html_parts.append(f'<h1 id="{anchor}">{line}</h1>')
        elif line.startswith("§"):
            anchor = ensure_unique(_simple_slug(line.replace("§", "paragraf")))
            h2 = {"level": 2, "title": line, "anchor": anchor}
            if last_h1:
                last_h1.setdefault("children", []).append(h2)
            else:
                toc.append(h2)
            html_parts.append(f'<h2 id="{anchor}">{line}</h2>')
        else:
            html_parts.append(f"<p>{line}</p>")

    return toc, "\n".join(html_parts)


def add_expired_j_messages_section(body_html: str, replaces: List[str]) -> str:
    """
    Add "Utgåtte j-meldinger" section at the end of body_html if replaces list is not empty.
    This matches the format used in Fiskeridirektoratet documents.
    
    Args:
        body_html: Current HTML body content
        replaces: List of J-message IDs that this document replaces
    
    Returns:
        body_html with "Utgåtte j-meldinger" section appended if replaces is not empty
    """
    if not replaces or len(replaces) == 0:
        return body_html
    
    # Filter out empty strings and None values
    valid_replaces = [r for r in replaces if r and str(r).strip()]
    if len(valid_replaces) == 0:
        return body_html
    
    # Create the section HTML (similar to Fiskeridirektoratet format)
    section_html = '\n<div style="margin-top: 32px; padding-top: 20px; border-top: 1px solid #d1d5db;">'
    section_html += '\n<h2 style="color: #1e3a8a; font-weight: bold; font-size: 1.25em; margin-bottom: 16px; margin-top: 0;">Utgåtte j-meldinger</h2>'
    section_html += '\n<div style="margin-left: 8px; line-height: 1.8;">'
    
    # Add each J-ID as a link-like element (underlined, blue, similar to original)
    for idx, j_id in enumerate(valid_replaces):
        j_id_str = str(j_id).strip()
        if j_id_str:
            # Add comma separator except for last item
            separator = ', ' if idx < len(valid_replaces) - 1 else ''
            section_html += f'<span style="color: #2563eb; text-decoration: underline; cursor: pointer;">{j_id_str}</span>{separator}'
    
    section_html += '\n</div>'
    section_html += '\n</div>'
    
    # Append to body_html
    return body_html + section_html


def extract_json_from_llm_response(response: str) -> dict:
    """
    Robust JSON parser for LLM responses.
    Handles responses with extra text, XML tags, markdown, etc.
    
    Args:
        response: Raw LLM response that should contain JSON
        
    Returns:
        Parsed JSON dict, or empty dict if parsing fails
    """
    # Check for empty or very short responses
    if not response or len(response.strip()) < 10:
        print(f"[JSON_PARSER] ⚠️ Response is empty or too short (len={len(response)})")
        print(f"[JSON_PARSER] Raw response: '{response}'")
        return {}
    
    try:
        json_str = response.strip()
        original_length = len(json_str)
        
        # Remove XML-like tags that some models add (e.g., <think>, <answer>)
        json_str = re.sub(r'<[^>]+>', '', json_str)
        
        # Check if response was ONLY tags (nothing left after removing them)
        if len(json_str.strip()) < 5:
            print(f"[JSON_PARSER] ⚠️ Response contained only XML tags, no actual content")
            print(f"[JSON_PARSER] Original ({original_length} chars): {response[:500]}")
            return {}
        
        # Remove markdown code blocks if present
        if "```" in json_str:
            match = re.search(r'```(?:json)?\s*(.*?)\s*```', json_str, re.DOTALL)
            if match:
                json_str = match.group(1).strip()
            else:
                json_str = re.sub(r'```[^`]*```', '', json_str).strip()
        
        # Extract content between first '{' and last '}'
        first_brace = json_str.find('{')
        last_brace = json_str.rfind('}')
        if first_brace != -1 and last_brace != -1:
            json_str = json_str[first_brace:last_brace + 1]
        else:
            # No braces found - not valid JSON
            print(f"[JSON_PARSER] ⚠️ No JSON object found in response")
            print(f"[JSON_PARSER] Cleaned text: {json_str[:500]}...")
            return {}
        
        # Try to find JSON object if still not starting with '{'
        if not json_str.startswith('{'):
            json_match = re.search(r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}', json_str, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
        
        parsed = json.loads(json_str)
        if isinstance(parsed, dict):
            return parsed
        else:
            print(f"[JSON_PARSER] ⚠️ Parsed value is not a dict: {type(parsed)}")
            return {}
            
    except json.JSONDecodeError as e:
        print(f"[JSON_PARSER] ❌ Failed to parse JSON: {e}")
        print(f"[JSON_PARSER] Response length: {len(response)} chars")
        print(f"[JSON_PARSER] First 500 chars: {response[:500]}")
        print(f"[JSON_PARSER] Last 200 chars: {response[-200:]}")
        return {}
    except Exception as e:
        print(f"[JSON_PARSER] ❌ Unexpected error: {e}")
        print(f"[JSON_PARSER] Response: {response[:500]}")
        return {}


def build_metadata_prompt(header_text: str, body_text: str) -> str:
    """
    Build metadata extraction prompt using versioned template.
    
    Args:
        header_text: Document header text
        body_text: Document body text (will be truncated to 4000 chars)
    
    Returns:
        Formatted prompt string
    """
    # Limit body text to 4000 characters to avoid token limits
    body_text_limited = body_text[:4000]
    
    # Try to load from versioned template
    if load_prompt_template:
        try:
            template = load_prompt_template("metadata")
            # Format template with actual values
            return template.format(header_text=header_text, body_text=body_text_limited)
        except Exception as e:
            print(f"[J-MESSAGES] ⚠️ Failed to load versioned prompt, using fallback: {e}")
    
    # Fallback to hardcoded prompt (original implementation)
    return f"""Extract metadata from this Norwegian fishing regulation document.
Return ONLY valid JSON. No explanations, no thinking, just JSON.

Required JSON format:
{{
  "j_id": "J-XXX-YYYY",
  "title": "document title",
  "replaces_id": ["J-XXX-YYYY", "J-XXX-YYYY"] or [],
  "replaced_by_id": "J-XXX-YYYY or null",
  "status": "active/inactive/repealed",
  "valid_from": "YYYY-MM-DD or null",
  "valid_to": "YYYY-MM-DD or null",
  "category": ["Annet", "Bunnfisk", "Pelagisk fisk"] or [],
  "area": ["Andre lands soner", "Internasjonal farvann", "Nord for 62° N", "Sør for 62° N"] or []
}}

IMPORTANT RULES: 
- "replaces_id" must be an ARRAY of J-message IDs that this document replaces (Norwegian UI label: "Erstatter"). Look for patterns like "(J-XXX-YYYY, J-XXX-YYYY UTGÅR)" in the document header and extract ALL J-message IDs mentioned. If no J-messages are replaced, use an empty array [].
- "replaced_by_id" is the J-message that replaces THIS document (Norwegian UI label: "Erstattet av").
- "category" must be an ARRAY containing one or more of these three values: "Annet", "Bunnfisk", or "Pelagisk fisk". A document can have multiple categories.
  * Include "Bunnfisk" if the document mentions bottom-dwelling fish species (e.g., torsk/cod, hyse/haddock, sei/saithe, blåkveite/blue halibut, rognkjeks/lumpfish, kongekrabbe/king crab, etc.)
  * Include "Pelagisk fisk" if the document mentions pelagic fish species (e.g., makrell/mackerel, sild/herring, brisling/sprat, etc.)
  * Include "Annet" if the document does not clearly mention Bunnfisk or Pelagisk fisk species, or if it's a general regulation
  * A document can have multiple categories (e.g., ["Bunnfisk", "Pelagisk fisk"] if it mentions both types)
  * If no specific category is identified, use ["Annet"] (array with one element, not empty array)
- "area" must be an ARRAY containing one or more of these four values: "Andre lands soner", "Internasjonal farvann", "Nord for 62° N", "Sør for 62° N". A document can have multiple areas. If no area is found, use an empty array [].
- You MUST include both "category" and "area" fields in your JSON response. Category must always be an array with at least one value (use ["Annet"] if no specific category is found).
- Look for geographical references in the document to determine the areas (e.g., "nord for 62°", "sør for 62°", "internasjonalt", etc.). A document may mention multiple areas.

Document text:
\"\"\"{header_text}\n\n{body_text_limited}\"\"\"

JSON:"""


def analyze_text_content(
    text_content: str,
    request_headers: Dict[str, str] = None,
    complexity: str = None,
    temperature: float = None
) -> Dict[str, Any]:
    """
    Helper function to analyze J-melding text content directly (for evaluator use).
    
    Args:
        text_content: Full text of the J-melding document
        request_headers: Optional dict of request headers for API config
        complexity: Optional AI complexity (low/medium/high); default "low"
        temperature: Optional sampling temperature (0.0-2.0); default None
    
    Returns:
        Dict with metadata, toc, body_html, and raw_text
    """
    # Split text into lines
    paragraphs = [ln for ln in text_content.split("\n") if ln.strip()]
    
    # Split header/body
    header_text, body_text = split_header_body(paragraphs)
    
    # Build TOC and HTML
    toc, body_html = build_toc_and_body_html(body_text)
    
    # Initialize metadata
    metadata: Dict[str, Any] = {
        "j_id": None,
        "title": None,
        "replaces_id": [],
        "replaced_by_id": None,
        "status": None,
        "valid_from": None,
        "valid_to": None,
        "category": ["Annet"],  # Default category as array
        "area": []
    }
    
    # Extract metadata using LLM if available
    if ask_ai_unified_sync:
        try:
            complexity_level = complexity if complexity in ("low", "medium", "high") else "low"
            temperature_value = temperature if temperature is None or (isinstance(temperature, (int, float)) and 0.0 <= float(temperature) <= 2.0) else None
            prompt = build_metadata_prompt(header_text, body_text)
            response = ask_ai_unified_sync(
                prompt=prompt + "\nGi svaret som STRICT JSON.",
                task_type="extraction",
                complexity=complexity_level,
                max_tokens=600,
                messages=None,
                request_headers=request_headers or {},
                temperature=temperature_value
            )
            
            # Parse JSON using the robust shared helper
            parsed = extract_json_from_llm_response(response)
            if parsed:
                metadata.update(parsed)
                print(f"[ANALYZER] Successfully extracted metadata: {list(parsed.keys())}")
            else:
                print(f"[ANALYZER] Failed to extract JSON from LLM response")
        except Exception as e:
            print(f"[ANALYZER] Failed to extract metadata with LLM: {e}")
    
    # Extract replaces list for adding "Utgåtte j-meldinger" section
    replaces_list = metadata.get("replaces_id") if isinstance(metadata.get("replaces_id"), list) else ([metadata.get("replaces_id")] if metadata.get("replaces_id") else [])
    
    # Add "Utgåtte j-meldinger" section at the end of body_html if there are replaces
    body_html_with_expired = add_expired_j_messages_section(body_html, replaces_list)
    
    return {
        "metadata": metadata,
        "toc": toc,
        "body_html": body_html_with_expired,
        "raw_text": body_text
    }


@router.post("/analyze")
async def analyze_j_message(
    request: Request,
    file: UploadFile = File(...),
    summary_length: str = Query(None, description="short|medium|long to include an AI summary"),
    complexity: str = Query("low", description="AI complexity level: low, medium, or high"),
    temperature: float = Query(None, description="Sampling temperature (0.0-2.0). Lower=more deterministic, higher=more creative."),
    metadata_prompt_override: Optional[str] = Form(None, description="Optional custom prompt for metadata extraction (e.g. from Prompt Manager Test). If it contains {header_text} or {body_text}, those are replaced; otherwise the document text is appended.")
):
    """
    Accept a .docx J-melding and return structured JSON with:
    - metadata (LLM-extracted)
    - toc (from detected headings)
    - body_html (with ids matching toc anchors)
    - raw_text (body only)
    """
    # Log incoming request info (for MCP debugging)
    request_headers_dict = dict(request.headers) if request else {}
    api_provider_header = request_headers_dict.get("x-api-provider", "NOT SET")
    itemai_url_header = request_headers_dict.get("x-itemai-url", "NOT SET")
    print(f"[J-MESSAGES ANALYZE] Request received")
    print(f"[J-MESSAGES ANALYZE]   x-api-provider: {api_provider_header}")
    print(f"[J-MESSAGES ANALYZE]   x-itemai-url: {itemai_url_header}")
    print(f"[J-MESSAGES ANALYZE]   All headers: {list(request_headers_dict.keys())}")
    
    file_bytes = await file.read()
    filename_lower = (file.filename or "").lower()

    # Read input (DOCX, PDF or JSON)
    paragraphs: List[str] = []
    header_text = ""
    body_text = ""
    toc: List[Dict[str, Any]] = []
    body_html = ""

    try:
        if filename_lower.endswith(".json"):
            data = _parse_json_relaxed(file_bytes.decode("utf-8"))
            full_text = _extract_text_from_json(data).strip()
            has_full = isinstance(data.get("body_html"), str) and (
                data.get("toc") is not None or data.get("id") or data.get("j_id") or data.get("title")
            )
            if has_full:
                # Use JSON as full analysis result
                j_id = data.get("id") or data.get("j_id")
                title = data.get("title")
                replaces_id = data.get("replaces") or data.get("replaces_id") or []
                replaces_list = replaces_id if isinstance(replaces_id, list) else ([replaces_id] if replaces_id else [])
                replaced_by_id = data.get("replaced_by") or data.get("replaced_by_id")
                cat = data.get("category") or data.get("categories")
                category = cat if isinstance(cat, list) else ([cat] if cat else ["Annet"])
                ar = data.get("area") or []
                area = ar if isinstance(ar, list) else ([ar] if ar else [])
                toc = data.get("toc") or []
                body_html = data.get("body_html") or ""
                body_text = _extract_text_from_json(data).strip()
                summary_text = (data.get("summary") or "").strip()
                if summary_length and not summary_text and body_text and ask_ai_unified_sync:
                    try:
                        sum_prompt = f"""
Lag en {summary_length} oppsummering av forskriften nedenfor.
Returner ren tekst, uten markers eller Markdown.
Tekst:
\"\"\"{body_text[:12000]}\"\"\"
"""
                        complexity_level = complexity if complexity in ["low", "medium", "high"] else "low"
                        temperature_value = temperature if temperature is None or (0.0 <= float(temperature) <= 2.0) else None
                        summary_text = ask_ai_unified_sync(
                            prompt=sum_prompt, task_type="summarization", complexity=complexity_level,
                            max_tokens=700, messages=None, request_headers=dict(request.headers), temperature=temperature_value
                        ) or ""
                    except Exception:
                        summary_text = ""
                body_html_with_expired = add_expired_j_messages_section(body_html, replaces_list)
                return {
                    "id": j_id, "title": title, "status": data.get("status"), "valid_from": data.get("valid_from"), "valid_to": data.get("valid_to"),
                    "replaces": replaces_list, "replaced_by": replaced_by_id, "category": category, "area": area, "toc": toc, "body_html": body_html_with_expired,
                    "raw_text": body_text, "summary": summary_text, "summary_length": summary_length, "prompt_version": None, "debug": {"header_text": ""}
                }
            if not full_text:
                raise HTTPException(status_code=400, detail="JSON must contain 'raw_text' or 'body'")
            paragraphs = [ln for ln in full_text.split("\n") if ln.strip()]
        elif filename_lower.endswith(".docx"):
            paragraphs = read_docx_paragraphs(file_bytes)
            full_text = "\n".join(paragraphs)
        elif filename_lower.endswith(".pdf"):
            full_text = read_pdf_text(file_bytes)
            paragraphs = [ln for ln in full_text.split("\n") if ln.strip()]
        elif filename_lower.endswith(".html") or filename_lower.endswith(".htm"):
            html_str = file_bytes.decode("utf-8", errors="replace")
            for tag in ("</p>", "</div>", "</tr>", "</li>", "</h1>", "</h2>", "</h3>", "</h4>", "</h5>", "</h6>", "<br>", "<br/>", "<br />"):
                html_str = html_str.replace(tag, "\n")
            full_text = re.sub(r"<[^>]+>", " ", html_str)
            full_text = re.sub(r"[ \t]+", " ", full_text)
            full_text = re.sub(r"\n\s*\n", "\n", full_text).strip()
            paragraphs = [ln.strip() for ln in full_text.split("\n") if ln.strip()]
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use .docx, .pdf, .json or .html")
    except HTTPException:
        raise
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse file: {e}")

    header_text, body_text = split_header_body(paragraphs)
    toc, body_html = build_toc_and_body_html(body_text)

    metadata: Dict[str, Any] = {
        "j_id": None,
        "title": None,
        "replaces_id": [],
        "replaced_by_id": None,
        "status": None,
        "valid_from": None,
        "valid_to": None,
        "category": ["Annet"],  # Default category as array
        "area": []
    }

    summary_text: str = ""

            # Prefer unified LLM pipeline if available; otherwise, return without metadata
    try:
        if ask_ai_unified_sync:
            # Use custom prompt from Prompt Manager "Test" when provided; otherwise versioned/metadata template
            plain_text_prompt = False
            if metadata_prompt_override and (o := metadata_prompt_override.strip()):
                body_lim = body_text[:4000]
                if "{header_text}" in o or "{body_text}" in o:
                    prompt = o.replace("{body_text}", body_lim).replace("{header_text}", header_text)
                else:
                    prompt = o + '\n\nDocument text:\n"""' + header_text + '\n\n' + body_lim + '"""'
                _lower = o.lower()
                plain_text_prompt = any(
                    x in _lower for x in ["uten json", "without json", "kun med", "ren tekst", "rensede teksten", "uten markdown", "no json", "plain text only"]
                )
                if not plain_text_prompt:
                    prompt += "\n\nJSON:"
            else:
                prompt = build_metadata_prompt(header_text, body_text)
            # Pass request headers to allow API config from frontend/MCP
            request_headers_dict = dict(request.headers) if request else {}
            
            # Debug: log API config headers (check for both x-itemai-url and x-itemai-key)
            has_api_config = any(
                k in request_headers_dict 
                for k in ["x-api-provider", "x-openai-key", "x-openrouter-key", "x-itemai-url", "x-itemai-key"]
            )
            
            if has_api_config:
                provider = request_headers_dict.get("x-api-provider", "unknown")
                print(f"[J-MESSAGES] ✅ API config headers found: provider={provider}")
                if provider == "itemai":
                    itemai_url = request_headers_dict.get("x-itemai-url", "not set")
                    print(f"[J-MESSAGES]    → ItemAI URL: {itemai_url}")
                print(f"[J-MESSAGES]    → All headers: {list(request_headers_dict.keys())}")
            else:
                print("[J-MESSAGES] ⚠️ No API config headers found, will use .env fallback")
                print(f"[J-MESSAGES]    → Available headers: {list(request_headers_dict.keys())}")
            
            # Validate and use complexity from query parameter
            complexity_level = complexity if complexity in ["low", "medium", "high"] else "low"
            # Validate temperature (optional)
            temperature_value = temperature if temperature is None or (0.0 <= float(temperature) <= 2.0) else None
            
            if plain_text_prompt:
                final_prompt = prompt
                max_tokens_llm = 8000
            else:
                final_prompt = prompt + "\nGi svaret som STRICT JSON."
                max_tokens_llm = 600
            response = ask_ai_unified_sync(
                prompt=final_prompt,
                task_type="extraction",
                complexity=complexity_level,
                max_tokens=max_tokens_llm,
                messages=None,
                request_headers=request_headers_dict,
                temperature=temperature_value
            )
            if plain_text_prompt and response:
                # User asked for plain text only: use LLM response as body, no JSON
                body_text = (response or "").strip()
                # Format plain text as HTML (paragraphs, escape entities)
                import html
                body_html = "\n".join(
                    f"<p>{html.escape(ln)}</p>" for ln in body_text.split("\n") if ln.strip()
                ) if body_text else ""
                toc = []
                parsed = None
            else:
                parsed = extract_json_from_llm_response(response)
            if parsed:
                # Backward/alternate key support for replaced_by_id
                if "replaced_by" in parsed and "replaced_by_id" not in parsed:
                    parsed["replaced_by_id"] = parsed.get("replaced_by")
                if "replacedBy" in parsed and "replaced_by_id" not in parsed:
                    parsed["replaced_by_id"] = parsed.get("replacedBy")

                # Handle replaces_id: convert string to array (backward compatibility)
                if "replaces_id" in parsed:
                    replaces_id = parsed.get("replaces_id")
                    if isinstance(replaces_id, str) and replaces_id:
                        # Parse comma-separated string into array
                        replaces_array = [r.strip() for r in replaces_id.split(',') if r.strip()]
                        parsed["replaces_id"] = replaces_array
                    elif not isinstance(replaces_id, list):
                        parsed["replaces_id"] = []
                else:
                    parsed["replaces_id"] = []

                # Handle legacy formats: convert string or old "categories" field to array
                if "categories" in parsed and "category" not in parsed:
                    categories = parsed.get("categories", [])
                    if isinstance(categories, list):
                        parsed["category"] = categories
                    elif isinstance(categories, str):
                        parsed["category"] = [categories] if categories else ["Annet"]
                    else:
                        parsed["category"] = ["Annet"]
                    # Remove old field
                    parsed.pop("categories", None)
                
                # Convert string category to array (backward compatibility)
                if "category" in parsed and isinstance(parsed.get("category"), str):
                    cat_str = parsed.get("category")
                    if cat_str in ["Annet", "Bunnfisk", "Pelagisk fisk"]:
                        parsed["category"] = [cat_str]
                    else:
                        parsed["category"] = ["Annet"]
                
                # Ensure category is always an array with at least one valid value
                category = parsed.get("category", [])
                if not isinstance(category, list):
                    category = [category] if category else ["Annet"]
                
                # Filter out invalid categories and ensure at least one valid category
                valid_categories = ["Annet", "Bunnfisk", "Pelagisk fisk"]
                original_category = list(category)
                category = [c for c in category if c in valid_categories]
                if not category:
                    category = ["Annet"]
                    print(f"[J-MESSAGES] ⚠️ Category was invalid (got {original_category!r}), defaulting to ['Annet']")

                parsed["category"] = category
                
                metadata.update(parsed)
                print(f"[J-MESSAGES] ✅ Successfully extracted metadata: {list(parsed.keys())}")
                print(f"[J-MESSAGES] Metadata values: id={parsed.get('j_id')}, title={parsed.get('title')[:50] if parsed.get('title') else None}, category={parsed.get('category')}, area={parsed.get('area')}")
            else:
                if not plain_text_prompt:
                    print(f"[J-MESSAGES] ⚠️ Failed to extract JSON from LLM response")
                # Leave defaults if parsing fails; front-end can still render (or plain-text result already set)
            # Optional summary (skip when Test returned plain text only)
            if summary_length and not plain_text_prompt:
                sum_prompt = f"""
Lag en {summary_length} oppsummering av forskriften nedenfor. 
Returner ren tekst, uten markers eller Markdown.
Tekst:
\"\"\"{body_text[:12000]}\"\"\"
"""
                try:
                    # Use same complexity level for summary
                    summary_text = ask_ai_unified_sync(
                        prompt=sum_prompt,
                        task_type="summarization",
                        complexity=complexity_level,
                        max_tokens=700,
                        messages=None,
                        request_headers=request_headers_dict,
                        temperature=temperature_value
                    ) or ""
                except Exception as _:
                    summary_text = ""
    except Exception as e:
        # Non-fatal; proceed with minimal payload
        print(f"[J-MESSAGES] Metadata extraction failed: {e}")

    # Get prompt version for audit trail
    prompt_version = None
    if get_active_version:
        try:
            prompt_version = get_active_version()
        except Exception:
            pass
    
    # Extract replaces list for adding "Utgåtte j-meldinger" section
    replaces_list = metadata.get("replaces_id") if isinstance(metadata.get("replaces_id"), list) else ([metadata.get("replaces_id")] if metadata.get("replaces_id") else [])
    
    # Add "Utgåtte j-meldinger" section at the end of body_html if there are replaces
    body_html_with_expired = add_expired_j_messages_section(body_html, replaces_list)
    
    result = {
        "id": metadata.get("j_id"),
        "title": metadata.get("title"),
        "status": metadata.get("status"),
        "valid_from": metadata.get("valid_from"),
        "valid_to": metadata.get("valid_to"),
        "replaces": replaces_list,
        "replaced_by": metadata.get("replaced_by_id") or metadata.get("replaced_by"),
        "category": _normalize_category(metadata.get("category"), metadata.get("categories")),
        "area": metadata.get("area") if isinstance(metadata.get("area"), list) else ([metadata.get("area")] if metadata.get("area") else []),
        "toc": toc,
        "body_html": body_html_with_expired,
        "raw_text": body_text,
        "summary": summary_text,
        "summary_length": summary_length,
        "prompt_version": prompt_version,  # Version of prompt used for this analysis
        "debug": {
            "header_text": header_text
        }
    }
    return result


@router.get("/prompt")
async def get_native_prompt(prompt_type: str = Query("metadata", description="Type of prompt: 'metadata' or 'note'")):
    """
    Get the current native prompt template.
    Used by frontend to display the "Native Prompt" in the UI.
    
    Args:
        prompt_type: "metadata" or "note" (default: "metadata")
    
    Returns:
        Dict with:
        - version: Active prompt version (e.g., "v1.0.0")
        - template: Prompt template string (with placeholders like {header_text}, {body_text})
        - type: Prompt type ("metadata" or "note")
    """
    try:
        if get_active_version:
            version = get_active_version()
        else:
            version = "unknown"
        
        if load_prompt_template:
            try:
                template = load_prompt_template(prompt_type)
                return {
                    "version": version,
                    "template": template,
                    "type": prompt_type
                }
            except FileNotFoundError as e:
                raise HTTPException(status_code=404, detail=f"Prompt template not found: {e}")
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to load prompt: {e}")
        else:
            # Fallback: return hardcoded prompt info
            if prompt_type == "metadata":
                # Return a simplified version of the hardcoded prompt
                template = """Extract metadata from this Norwegian fishing regulation document.
Return ONLY valid JSON. No explanations, no thinking, just JSON.

Required JSON format:
{{
  "j_id": "J-XXX-YYYY",
  "title": "document title",
  "replaces_id": ["J-XXX-YYYY", "J-XXX-YYYY"] or [],
  "replaced_by_id": "J-XXX-YYYY or null",
  "status": "active/inactive/repealed",
  "valid_from": "YYYY-MM-DD or null",
  "valid_to": "YYYY-MM-DD or null",
  "category": ["Annet", "Bunnfisk", "Pelagisk fisk"] or [],
  "area": ["Andre lands soner", "Internasjonal farvann", "Nord for 62° N", "Sør for 62° N"] or []
}}

IMPORTANT RULES: 
- "replaces_id" must be an ARRAY of J-message IDs that this document replaces (Norwegian UI label: "Erstatter"). Look for patterns like "(J-XXX-YYYY, J-XXX-YYYY UTGÅR)" in the document header and extract ALL J-message IDs mentioned. If no J-messages are replaced, use an empty array [].
- "replaced_by_id" is the J-message that replaces THIS document (Norwegian UI label: "Erstattet av").
- "category" must be an ARRAY containing one or more of these three values: "Annet", "Bunnfisk", or "Pelagisk fisk". A document can have multiple categories.
  * Include "Bunnfisk" if the document mentions bottom-dwelling fish species (e.g., torsk/cod, hyse/haddock, sei/saithe, blåkveite/blue halibut, rognkjeks/lumpfish, kongekrabbe/king crab, etc.)
  * Include "Pelagisk fisk" if the document mentions pelagic fish species (e.g., makrell/mackerel, sild/herring, brisling/sprat, etc.)
  * Include "Annet" if the document does not clearly mention Bunnfisk or Pelagisk fisk species, or if it's a general regulation
  * A document can have multiple categories (e.g., ["Bunnfisk", "Pelagisk fisk"] if it mentions both types)
  * If no specific category is identified, use ["Annet"] (array with one element, not empty array)
- "area" must be an ARRAY containing one or more of these four values: "Andre lands soner", "Internasjonal farvann", "Nord for 62° N", "Sør for 62° N". A document can have multiple areas. If no area is found, use an empty array [].
- You MUST include both "category" and "area" fields in your JSON response. Category must always be an array with at least one value (use ["Annet"] if no specific category is found).
- Look for geographical references in the document to determine the areas (e.g., "nord for 62°", "sør for 62°", "internasjonalt", etc.). A document may mention multiple areas.

Document text:
\"\"\"{header_text}\n\n{body_text}\"\"\"

JSON:"""
            elif prompt_type == "note":
                template = """Extract note metadata from this Norwegian J-melding addendum document.
Return ONLY valid JSON. No explanations, no thinking, just JSON.

Required JSON format:
{{
  "target_j_id": "J-XXX-YYYY",
  "note_type": "addendum",
  "valid_from": "YYYY-MM-DD",
  "valid_to": "YYYY-MM-DD",
  "affected_sections": ["Chapter 1", "§ 7"],
  "actions": ["amend", "replace"],
  "summary": "Brief change description"
}}

Document text:
\"\"\"{body_text}\"\"\"

JSON:"""
            else:
                raise HTTPException(status_code=400, detail=f"Unknown prompt type: {prompt_type}")
            
            return {
                "version": "fallback",
                "template": template,
                "type": prompt_type
            }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Unexpected error: {e}")


@router.post("/save")
async def save_j_message(data: Dict[str, Any] = Body(...)):
    """
    Save analyzed J-message to MongoDB.
    Expects JSON payload similar to /analyze result and optional filename.
    """
    if j_messages_collection is None:
        raise HTTPException(status_code=500, detail="MongoDB not configured")
    try:
        # header_text is needed so Evaluate (pairs Library) can run on full document = header + body
        header_text = (data.get("debug") or {}).get("header_text") or ""
        doc: Dict[str, Any] = {
            "title": data.get("title"),
            "j_id": data.get("id"),
            "status": data.get("status"),
            "valid_from": data.get("valid_from"),
            "valid_to": data.get("valid_to"),
            "replaces": data.get("replaces") if isinstance(data.get("replaces"), list) else ([data.get("replaces")] if data.get("replaces") else []),
            "replaced_by": data.get("replaced_by"),
            "category": _normalize_category(data.get("category"), data.get("categories")),
            "area": data.get("area") if isinstance(data.get("area"), list) else ([data.get("area")] if data.get("area") else []),
            "toc": data.get("toc") or [],
            "body_html": data.get("body_html") or "",
            "summary": data.get("summary") or "",
            "summary_length": data.get("summary_length"),
            "raw_text": data.get("raw_text") or "",
            "header_text": header_text,
            "filename": data.get("filename") or "",
            "prompt_version": data.get("prompt_version"),  # Store prompt version for audit
            "created_at": datetime.utcnow(),
            "module": "j_messages"
        }
        result = await j_messages_collection.insert_one(doc)
        return {"success": True, "id": str(result.inserted_id)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Save failed: {e}")


@router.get("/list")
async def list_j_messages(
    status: str = Query(None, description="Filter by status (e.g., 'Gjeldende', 'Utgått')"),
    category: str = Query(None, description="Filter by category"),
    search: str = Query(None, description="Search in title or content")
):
    """
    List saved J-messages from MongoDB.
    Supports MCP tool: list_j_meldinger
    """
    if j_messages_collection is None:
        return {"success": True, "items": [], "total": 0}
    
    # Build query filter
    query: Dict[str, Any] = {}
    if status:
        query["status"] = status
    or_clauses = []
    if category:
        # Support both old (string) and new (array) format
        or_clauses.append({"$or": [
            {"category": category},           # Exact match for string
            {"category": {"$in": [category]}},  # Array contains this value
            {"categories": {"$in": [category]}}  # Legacy format
        ]})
    if search:
        or_clauses.append({"$or": [
            {"title": {"$regex": search, "$options": "i"}},
            {"j_id": {"$regex": search, "$options": "i"}},
            {"raw_text": {"$regex": search, "$options": "i"}}
        ]})
    if len(or_clauses) == 1:
        query.update(or_clauses[0])   # Single filter: apply $or directly
    elif len(or_clauses) > 1:
        query["$and"] = or_clauses    # Both filters: combine with $and
    
    items: List[Dict[str, Any]] = []
    async for doc in j_messages_collection.find(query).sort("created_at", -1):
        doc["id"] = str(doc.pop("_id"))
        items.append(doc)
    return {"success": True, "items": items, "total": len(items)}


@router.put("/update/{doc_id}")
async def update_j_message(doc_id: str, data: Dict[str, Any] = Body(...)):
    """
    Update an existing J-message in MongoDB.
    """
    if j_messages_collection is None:
        raise HTTPException(status_code=500, detail="MongoDB not configured")
    try:
        from bson import ObjectId
        update_doc: Dict[str, Any] = {
            "updated_at": datetime.utcnow()
        }
        # Only update fields that are provided
        if "title" in data:
            update_doc["title"] = data.get("title")
        if "j_id" in data:
            update_doc["j_id"] = data.get("j_id")
        if "status" in data:
            update_doc["status"] = data.get("status")
        if "valid_from" in data:
            update_doc["valid_from"] = data.get("valid_from")
        if "valid_to" in data:
            update_doc["valid_to"] = data.get("valid_to")
        if "replaces" in data:
            replaces_value = data.get("replaces")
            # Handle replaces as array
            if isinstance(replaces_value, list):
                # Filter out empty strings and ensure valid J-ID format
                replaces_value = [r.strip() for r in replaces_value if r and r.strip()]
                update_doc["replaces"] = replaces_value if replaces_value else []
            elif isinstance(replaces_value, str):
                # Parse comma-separated string into array
                if replaces_value:
                    replaces_array = [r.strip() for r in replaces_value.split(',') if r.strip()]
                    update_doc["replaces"] = replaces_array
                else:
                    update_doc["replaces"] = []
            else:
                update_doc["replaces"] = []
        if "replaced_by" in data:
            update_doc["replaced_by"] = data.get("replaced_by")
        # Track fields to unset separately
        unset_fields = {}
        
        if "category" in data:
            category_value = data.get("category")
            # Handle category as array
            if isinstance(category_value, list):
                # Ensure it's a valid array (filter invalid values)
                valid_categories = ["Annet", "Bunnfisk", "Pelagisk fisk"]
                category_value = [c for c in category_value if c in valid_categories]
                if not category_value:
                    category_value = ["Annet"]
                update_doc["category"] = category_value
            elif category_value == "" or category_value is None:
                update_doc["category"] = ["Annet"]  # Default to array with Annet
            else:
                # Convert string to array
                if category_value in ["Annet", "Bunnfisk", "Pelagisk fisk"]:
                    update_doc["category"] = [category_value]
                else:
                    update_doc["category"] = ["Annet"]
            # Remove old categories field if it exists
            unset_fields["categories"] = ""
        elif "categories" in data:
            # Legacy support: convert old categories array to new category array
            cats = data.get("categories") or []
            if isinstance(cats, list):
                valid_categories = ["Annet", "Bunnfisk", "Pelagisk fisk"]
                cats = [c for c in cats if c in valid_categories]
                update_doc["category"] = cats if cats else ["Annet"]
            elif isinstance(cats, str):
                update_doc["category"] = [cats] if cats in ["Annet", "Bunnfisk", "Pelagisk fisk"] else ["Annet"]
            else:
                update_doc["category"] = ["Annet"]
            unset_fields["categories"] = ""
        if "area" in data:
            area_value = data.get("area")
            # Ensure area is always an array
            if isinstance(area_value, list):
                update_doc["area"] = area_value
            elif area_value:
                update_doc["area"] = [area_value]
            else:
                update_doc["area"] = []
        if "toc" in data:
            update_doc["toc"] = data.get("toc") or []
        if "body_html" in data:
            update_doc["body_html"] = data.get("body_html") or ""
        if "summary" in data:
            update_doc["summary"] = data.get("summary") or ""
        if "summary_length" in data:
            update_doc["summary_length"] = data.get("summary_length")
        if "raw_text" in data:
            update_doc["raw_text"] = data.get("raw_text") or ""
        if "filename" in data:
            update_doc["filename"] = data.get("filename") or ""
        
        # Build update operation with both $set and $unset if needed
        update_operation = {}
        if update_doc:
            update_operation["$set"] = update_doc
        if unset_fields:
            update_operation["$unset"] = unset_fields
        
        result = await j_messages_collection.update_one(
            {"_id": ObjectId(doc_id)},
            update_operation if update_operation else {"$set": {"updated_at": datetime.utcnow()}}
        )
        if result.matched_count == 0:
            raise HTTPException(status_code=404, detail="J-message not found")
        return {"success": True, "id": doc_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Update failed: {e}")


@router.delete("/delete/{doc_id}")
async def delete_j_message(doc_id: str):
    if j_messages_collection is None:
        raise HTTPException(status_code=500, detail="MongoDB not configured")
    try:
        from bson import ObjectId
        res = await j_messages_collection.delete_one({"_id": ObjectId(doc_id)})
        return {"success": res.deleted_count == 1}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Delete failed: {e}")


def build_note_prompt(body_text: str) -> str:
    """
    Build note extraction prompt using versioned template.
    
    Args:
        body_text: Document body text (will be truncated to 12000 chars)
    
    Returns:
        Formatted prompt string
    """
    # Limit body text to 12000 characters to avoid token limits
    body_text_limited = body_text[:12000]
    
    # Try to load from versioned template
    if load_prompt_template:
        try:
            template = load_prompt_template("note")
            # Format template with actual values
            return template.format(body_text=body_text_limited)
        except Exception as e:
            print(f"[J-MESSAGES] ⚠️ Failed to load versioned prompt, using fallback: {e}")
    
    # Fallback to hardcoded prompt (original implementation)
    return f"""Extract note metadata from this Norwegian J-melding addendum document.
Return ONLY valid JSON. No explanations, no thinking, just JSON.

Required JSON format:
{{
  "target_j_id": "J-XXX-YYYY",
  "note_type": "addendum",
  "valid_from": "YYYY-MM-DD",
  "valid_to": "YYYY-MM-DD",
  "affected_sections": ["Chapter 1", "§ 7"],
  "actions": ["amend", "replace"],
  "summary": "Brief change description"
}}

Document text:
\"\"\"{body_text_limited}\"\"\"

JSON:"""


@router.post("/analyze-note")
async def analyze_j_note(
    request: Request, 
    file: UploadFile = File(...),
    complexity: str = Query("low", description="AI complexity level: low, medium, or high"),
    temperature: float = Query(None, description="Sampling temperature (0.0-2.0). Lower=more deterministic, higher=more creative.")
):
    """
    Analyze a J-melding 'note' document (DOCX/PDF). Returns structured 'note' fields.
    """
    file_bytes = await file.read()
    filename_lower = (file.filename or "").lower()

    # Read text from file
    try:
        if filename_lower.endswith(".docx"):
            paragraphs = read_docx_paragraphs(file_bytes)
            full_text = "\n".join(paragraphs)
        elif filename_lower.endswith(".pdf"):
            full_text = read_pdf_text(file_bytes)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use .docx or .pdf")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse file: {e}")

    # Use LLM to extract note structure
    note_data: Dict[str, Any] = {
        "target_j_id": None,
        "note_type": None,
        "valid_from": None,
        "valid_to": None,
        "affected_sections": [],
        "actions": [],
        "summary": ""
    }
    try:
        if ask_ai_unified_sync:
            prompt = build_note_prompt(full_text)
            # Validate and use complexity from query parameter
            complexity_level = complexity if complexity in ["low", "medium", "high"] else "low"
            temperature_value = temperature if temperature is None or (0.0 <= float(temperature) <= 2.0) else None
            
            resp = ask_ai_unified_sync(
                prompt=prompt,
                task_type="extraction",
                complexity=complexity_level,
                max_tokens=600,
                messages=None,
                request_headers=dict(request.headers),
                temperature=temperature_value
            )
            # Parse JSON using robust helper function
            parsed = extract_json_from_llm_response(resp)
            if parsed:
                note_data.update(parsed)
                print(f"[J-MESSAGES NOTE] ✅ Successfully parsed note data")
            else:
                print(f"[J-MESSAGES NOTE] ⚠️ Failed to extract JSON from LLM response")
    except Exception as e:
        print(f"[J-MESSAGES NOTE] extraction failed: {e}")

    return {
        "type": "note",
        "note": note_data,
        "raw_text": full_text
    }


@router.post("/extract-text")
async def extract_text_only(
    file: UploadFile = File(...)
):
    """
    Extract raw text from a DOCX or PDF file without any analysis.
    Returns only the raw text content.
    """
    file_bytes = await file.read()
    filename_lower = (file.filename or "").lower()

    # Read input (DOCX, PDF or JSON) - only extract text, no analysis
    try:
        if filename_lower.endswith(".docx"):
            paragraphs = read_docx_paragraphs(file_bytes)
            full_text = "\n".join(paragraphs)
        elif filename_lower.endswith(".pdf"):
            full_text = read_pdf_text(file_bytes)
        elif filename_lower.endswith(".json"):
            data = _parse_json_relaxed(file_bytes.decode("utf-8"))
            full_text = _extract_text_from_json(data).strip()
            if not full_text:
                raise HTTPException(
                    status_code=400,
                    detail="JSON must contain 'raw_text', 'body', 'text' or Enonic 'data.blocks.text.text'"
                )
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use .docx, .pdf or .json")
    except HTTPException:
        raise
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse file: {e}")

    return {
        "raw_text": full_text,
        "filename": file.filename
    }


@router.post("/export-docx")
async def export_docx(data: Dict[str, Any] = Body(...)):
    """
    Build a simple .docx from the analyzed J-message.
    Uses title, metadata, optional summary and raw_text/body_html (as plain text).
    """
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx not available on server")
    try:
        # Prefer raw_text to avoid HTML parsing; fallback to stripped body_html
        body_text = data.get("raw_text") or ""
        if not body_text and data.get("body_html"):
            # Very naive HTML → text fallback
            body_text = re.sub("<[^>]+>", " ", data.get("body_html") or "")
            body_text = re.sub(r"\s+", " ", body_text).strip()

        doc = Document()
        title = data.get("title") or data.get("id") or "J-message"
        doc.add_heading(title, 0)

        # Metadata
        meta_lines: List[str] = []
        if data.get("id"): meta_lines.append(f"ID: {data.get('id')}")
        if data.get("status"): meta_lines.append(f"Status: {data.get('status')}")
        if data.get("valid_from"): meta_lines.append(f"Valid from: {data.get('valid_from')}")
        if data.get("valid_to"): meta_lines.append(f"Valid to: {data.get('valid_to')}")
        replaces = data.get("replaces")
        if isinstance(replaces, list) and len(replaces) > 0:
            meta_lines.append(f"Replaces: {', '.join(replaces)}")
        elif replaces:
            meta_lines.append(f"Replaces: {replaces}")
        if data.get("replaced_by"): meta_lines.append(f"Replaced by: {data.get('replaced_by')}")
        # Support both old (string) and new (array) format
        category = data.get("category")
        if isinstance(category, list):
            if len(category) > 0:
                meta_lines.append(f"Category: {', '.join(category)}")
        elif category:
            meta_lines.append(f"Category: {category}")
        else:
            # Legacy support
            cats = data.get("categories") or []
            if isinstance(cats, list) and len(cats) > 0:
                meta_lines.append(f"Category: {', '.join(cats)}")
            elif cats:
                meta_lines.append(f"Category: {cats}")
        area = data.get("area")
        if isinstance(area, list) and len(area) > 0:
            meta_lines.append(f"Area: {', '.join(area)}")
        elif area and not isinstance(area, list):
            meta_lines.append(f"Area: {area}")
        if meta_lines:
            p = doc.add_paragraph()
            for line in meta_lines:
                p.add_run(line).italic = True
                p.add_run("\n")

        # Summary
        if data.get("summary"):
            doc.add_heading("Executive Summary", level=1)
            for line in str(data.get("summary")).split("\n"):
                doc.add_paragraph(line)

        # Table of Contents (Content/Innhold)
        toc = data.get("toc") or []
        if toc:
            doc.add_heading("Content", level=1)  # "Innhold" in Norwegian, but using "Content" for consistency
            for item in toc:
                # Add main heading (level 1: Kapittel)
                if item.get("level") == 1:
                    # Add level 1 item with bullet
                    p = doc.add_paragraph(item.get("title", ""), style="List Bullet")
                    # Add children (level 2: § paragraphs) as sub-items with indentation
                    children = item.get("children", [])
                    for child in children:
                        if child.get("level") == 2:
                            # Add as indented sub-item
                            p_child = doc.add_paragraph()
                            # Add bullet with indentation (using left indent of 0.5 inches)
                            if Inches:
                                p_child.paragraph_format.left_indent = Inches(0.5)
                            p_child.add_run("• " + child.get("title", ""))
                elif item.get("level") == 2:
                    # Standalone level 2 item (no parent level 1)
                    doc.add_paragraph(item.get("title", ""), style="List Bullet")

        # Body - parse body_html to include tables
        doc.add_heading("Body", level=1)
        body_html = data.get("body_html") or ""
        if body_html:
            # Parse HTML and convert to DOCX, preserving tables
            from html.parser import HTMLParser
            
            # Split body_html into chunks (paragraphs, headings, tables)
            # Simple approach: extract tables first, then process remaining text
            table_pattern = r'<table[^>]*>.*?</table>'
            tables = re.findall(table_pattern, body_html, re.DOTALL | re.IGNORECASE)
            
            # Remove tables from HTML to get text content
            text_content = re.sub(table_pattern, '<TABLE_PLACEHOLDER>', body_html, flags=re.DOTALL | re.IGNORECASE)
            
            # Process text content (headings and paragraphs)
            lines = text_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check for placeholder
                if '<TABLE_PLACEHOLDER>' in line:
                    # Insert first available table
                    if tables:
                        table_html = tables.pop(0)
                        try:
                            # Convert HTML table to DOCX table
                            html_table_to_docx(table_html, doc)
                        except Exception as e:
                            print(f"[J-MESSAGES EXPORT] ⚠️ Failed to convert table: {e}")
                            doc.add_paragraph("[Table conversion failed]")
                    continue
                
                # Process headings
                if line.startswith('<h1'):
                    match = re.search(r'<h1[^>]*>(.*?)</h1>', line, re.IGNORECASE)
                    if match:
                        doc.add_heading(match.group(1), level=1)
                elif line.startswith('<h2'):
                    match = re.search(r'<h2[^>]*>(.*?)</h2>', line, re.IGNORECASE)
                    if match:
                        doc.add_heading(match.group(1), level=2)
                elif line.startswith('<p'):
                    # Extract text from paragraph, removing HTML tags
                    text = re.sub(r'<[^>]+>', '', line)
                    text = re.sub(r'\s+', ' ', text).strip()
                    if text:
                        doc.add_paragraph(text)
                else:
                    # Plain text line
                    text = re.sub(r'<[^>]+>', '', line)
                    text = re.sub(r'\s+', ' ', text).strip()
                    if text:
                        doc.add_paragraph(text)
            
            # Add any remaining tables
            for table_html in tables:
                try:
                    html_table_to_docx(table_html, doc)
                except Exception as e:
                    print(f"[J-MESSAGES EXPORT] ⚠️ Failed to convert table: {e}")
                    doc.add_paragraph("[Table conversion failed]")
        else:
            # Fallback to raw_text
            for line in (body_text or "").split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        # Stream as response
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        filename = f"{(data.get('id') or 'j-message').replace(' ', '_')}.docx"
        return StreamingResponse(
            bio,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {e}")


def load_lovteknikk_rules() -> str:
    """
    Load the lovteknikkboka-oppdatert.md document containing rules for writing j-meldinger.
    
    Returns:
        Content of the rules document as a string
    """
    import os
    # Get the project root directory
    current_file = __file__
    # __file__ is backend/routers/j_messages_analyzer.py
    # Go up 2 levels to get to project root
    project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(current_file))))
    rules_path = os.path.join(project_root, "docs-md", "lovteknikkboka-oppdatert.md")
    
    print(f"[J-MESSAGES PRE-ANALYZE] Looking for rules document at: {rules_path}")
    print(f"[J-MESSAGES PRE-ANALYZE] File exists: {os.path.exists(rules_path)}")
    
    try:
        with open(rules_path, "r", encoding="utf-8") as f:
            content = f.read()
            print(f"[J-MESSAGES PRE-ANALYZE] ✅ Rules document loaded successfully, {len(content)} characters")
            return content
    except Exception as e:
        print(f"[J-MESSAGES PRE-ANALYZE] ⚠️ Failed to load rules document: {e}")
        import traceback
        print(f"[J-MESSAGES PRE-ANALYZE] ⚠️ Traceback: {traceback.format_exc()}")
        return ""


def get_pre_analyze_prompt_template() -> str:
    """
    Get the template for the pre-analyze prompt (without document-specific content).
    This is used to display the prompt structure in the UI.
    
    Returns:
        Prompt template string with placeholders
    """
    return """Du er en ekspert på norsk lovteknikk og forskriftsarbeid. Du skal analysere og omskrive en J-melding fra Fiskeridirektoratet basert på reglene i "Lovteknikk og lovforberedelse - Veiledning om lov- og forskriftsarbeid".

VIKTIGE INSTRUKSJONER:
1. Les nøye gjennom reglene i lovteknikkboka nedenfor
2. Analyser den opprinnelige J-meldingen
3. Omskriv J-meldingen slik at den følger reglene så nøye som mulig
4. Behold all viktig informasjon fra originalen
5. Forbedre struktur, terminologi og språkføring i henhold til reglene
6. Sørg for at dokumentet følger norsk lovteknisk praksis

REGLER FRA LOVTEKNIKKBOKA:
[Reglene fra dokumentet lovteknikkboka-oppdatert.md settes inn her - begrenset til 20000 tegn]

OPPRINNELIG J-MELDING:
[Teksten fra det opprinnelige J-melding-dokumentet settes inn her - begrenset til 8000 tegn]

OPPGAVE:
Omskriv J-meldingen over slik at den følger reglene fra lovteknikkboka. Behold all viktig informasjon, men forbedre strukturen, terminologien og språkføringen. Returner dokumentet i samme format som originalen (med kapitler, paragrafer, etc.), men med forbedringer basert på reglene.

Returner kun det omskrevne dokumentet, uten ekstra kommentarer eller forklaringer."""


def build_pre_analyze_prompt(document_text: str, rules_text: str) -> str:
    """
    Build prompt for pre-analyzing a j-melding based on lovteknikk rules.
    
    Args:
        document_text: The original j-melding document text
        rules_text: The lovteknikk rules document text
    
    Returns:
        Formatted prompt string
    """
    # Limit document text to avoid token limits (keep it reasonable)
    document_text_limited = document_text[:8000]
    
    # Limit rules text to most relevant sections (first 20000 chars should cover key rules)
    rules_text_limited = rules_text[:20000]
    
    return f"""Du er en ekspert på norsk lovteknikk og forskriftsarbeid. Du skal analysere og omskrive en J-melding fra Fiskeridirektoratet basert på reglene i "Lovteknikk og lovforberedelse - Veiledning om lov- og forskriftsarbeid".

VIKTIGE INSTRUKSJONER:
1. Les nøye gjennom reglene i lovteknikkboka nedenfor
2. Analyser den opprinnelige J-meldingen
3. Omskriv J-meldingen slik at den følger reglene så nøye som mulig
4. Behold all viktig informasjon fra originalen
5. Forbedre struktur, terminologi og språkføring i henhold til reglene
6. Sørg for at dokumentet følger norsk lovteknisk praksis

REGLER FRA LOVTEKNIKKBOKA:
{rules_text_limited}

OPPRINNELIG J-MELDING:
{document_text_limited}

OPPGAVE:
Omskriv J-meldingen over slik at den følger reglene fra lovteknikkboka. Behold all viktig informasjon, men forbedre strukturen, terminologien og språkføringen. Returner dokumentet i samme format som originalen (med kapitler, paragrafer, etc.), men med forbedringer basert på reglene.

Returner kun det omskrevne dokumentet, uten ekstra kommentarer eller forklaringer."""


@router.get("/pre-analyze/prompt")
async def get_pre_analyze_prompt():
    """
    Get the pre-analyze prompt template.
    Used by frontend to display the prompt in the UI.
    
    Returns:
        Dict with:
        - template: The prompt template string
        - description: Description of what the prompt does
    """
    try:
        template = get_pre_analyze_prompt_template()
        return {
            "template": template,
            "description": "Prompt used to rewrite J-messages according to lovteknikk rules. The prompt includes instructions and references the lovteknikkboka-oppdatert.md document."
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to load prompt template: {e}")


@router.post("/pre-analyze")
async def pre_analyze_j_message(
    request: Request,
    file: UploadFile = File(...),
    complexity: str = Query("medium", description="AI complexity level: low, medium, or high"),
    temperature: float = Query(None, description="Sampling temperature (0.0-2.0). Lower=more deterministic, higher=more creative.")
):
    """
    Pre-analyze a J-melding by rewriting it according to lovteknikk rules.
    Uses the lovteknikkboka-oppdatert.md document as reference for rewriting rules.
    """
    print(f"[J-MESSAGES PRE-ANALYZE] ✅✅✅ ENDPOINT CALLED - File: {file.filename}, Complexity: {complexity}, Temperature: {temperature}")
    try:
        file_bytes = await file.read()
        filename_lower = (file.filename or "").lower()
        print(f"[J-MESSAGES PRE-ANALYZE] ✅ File read successfully, size: {len(file_bytes)} bytes")

        # Read input (DOCX, PDF, JSON or HTML)
        try:
            if filename_lower.endswith(".docx"):
                paragraphs = read_docx_paragraphs(file_bytes)
                full_text = "\n".join(paragraphs)
            elif filename_lower.endswith(".pdf"):
                full_text = read_pdf_text(file_bytes)
                paragraphs = [ln for ln in full_text.split("\n") if ln.strip()]
            elif filename_lower.endswith(".json"):
                data = _parse_json_relaxed(file_bytes.decode("utf-8"))
                full_text = _extract_text_from_json(data).strip()
                if not full_text:
                    raise HTTPException(status_code=400, detail="JSON must contain 'raw_text', 'body' or Enonic 'data.blocks.text.text'")
            elif filename_lower.endswith(".html") or filename_lower.endswith(".htm"):
                html_str = file_bytes.decode("utf-8", errors="replace")
                # Insert newlines after block elements to preserve structure
                for tag in ("</p>", "</div>", "</tr>", "</li>", "</h1>", "</h2>", "</h3>", "</h4>", "</h5>", "</h6>", "<br>", "<br/>", "<br />"):
                    html_str = html_str.replace(tag, "\n")
                full_text = re.sub(r"<[^>]+>", " ", html_str)
                full_text = re.sub(r"[ \t]+", " ", full_text)
                full_text = re.sub(r"\n\s*\n", "\n", full_text).strip()
                paragraphs = [ln.strip() for ln in full_text.split("\n") if ln.strip()]
            else:
                raise HTTPException(status_code=400, detail="Unsupported file type. Use .docx, .pdf, .json or .html")
        except HTTPException:
            raise
        except json.JSONDecodeError as e:
            raise HTTPException(status_code=400, detail=f"Invalid JSON: {e}")
        except Exception as e:
            import traceback
            print(f"[J-MESSAGES PRE-ANALYZE] ❌ Failed to parse file: {e}")
            print(f"[J-MESSAGES PRE-ANALYZE] ❌ Traceback:\n{traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"Failed to parse file: {e}")

        # Load rules document
        print(f"[J-MESSAGES PRE-ANALYZE] Loading rules document...")
        rules_text = load_lovteknikk_rules()
        if not rules_text:
            print(f"[J-MESSAGES PRE-ANALYZE] ❌ Failed to load rules document")
            raise HTTPException(status_code=500, detail="Failed to load lovteknikk rules document")
        print(f"[J-MESSAGES PRE-ANALYZE] ✅ Rules document loaded, size: {len(rules_text)} characters")

        # Build prompt
        prompt = build_pre_analyze_prompt(full_text, rules_text)

        # Use LLM to rewrite the document
        try:
            if not ask_ai_unified_sync:
                raise HTTPException(status_code=500, detail="LLM service not available")
            
            complexity_level = complexity if complexity in ["low", "medium", "high"] else "medium"
            temperature_value = temperature if temperature is None or (0.0 <= float(temperature) <= 2.0) else None
            
            request_headers_dict = dict(request.headers) if request else {}
            
            print(f"[J-MESSAGES PRE-ANALYZE] Calling LLM with complexity={complexity_level}, temperature={temperature_value}")
            rewritten_text = ask_ai_unified_sync(
                prompt=prompt,
                task_type="rewriting",
                complexity=complexity_level,
                max_tokens=8000,  # Higher limit for rewriting tasks
                messages=None,
                request_headers=request_headers_dict,
                temperature=temperature_value
            )
            
            if not rewritten_text or rewritten_text.strip().startswith("[MOCKED"):
                raise HTTPException(status_code=500, detail="LLM returned invalid response")
            
            # Process the rewritten text to build TOC and HTML
            rewritten_paragraphs = [ln for ln in rewritten_text.split("\n") if ln.strip()]
            toc, body_html = build_toc_and_body_html("\n".join(rewritten_paragraphs))
            
            # Extract basic metadata (title from first line or heading)
            title = None
            if rewritten_paragraphs:
                first_line = rewritten_paragraphs[0].strip()
                if first_line and len(first_line) < 200:
                    title = first_line
            
            print(f"[J-MESSAGES PRE-ANALYZE] ✅ Successfully processed, returning result")
            return {
                "success": True,
                "id": None,  # Pre-analysis doesn't create a saved document
                "title": title,
                "toc": toc,
                "body_html": body_html,
                "raw_text": rewritten_text,
                "type": "pre-analyzed"
            }
        except HTTPException:
            raise
        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            print(f"[J-MESSAGES PRE-ANALYZE] ❌ Error: {e}")
            print(f"[J-MESSAGES PRE-ANALYZE] ❌ Traceback:\n{error_trace}")
            raise HTTPException(status_code=500, detail=f"Pre-analysis failed: {str(e)}")
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"[J-MESSAGES PRE-ANALYZE] ❌ Unexpected error: {e}")
        print(f"[J-MESSAGES PRE-ANALYZE] ❌ Traceback:\n{error_trace}")
        raise HTTPException(status_code=500, detail=f"Pre-analysis failed: {str(e)}")


@router.post("/pre-analyze/export-docx")
async def export_pre_analyzed_docx(data: Dict[str, Any] = Body(...)):
    """
    Export a pre-analyzed J-message to DOCX format.
    Similar to regular export but without metadata extraction.
    """
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx not available on server")
    
    try:
        doc = Document()
        
        # Add title
        if data.get("title"):
            title_para = doc.add_heading(data["title"], level=1)
        
        # Add body content
        body_html = data.get("body_html", "")
        body_text = data.get("raw_text", "")
        
        if body_html:
            # Parse HTML and convert to DOCX
            lines = body_html.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check for headings
                if line.startswith("<h1") or line.startswith("<h2"):
                    # Extract text from heading
                    match = re.search(r">([^<]+)<", line)
                    if match:
                        level = 1 if "h1" in line else 2
                        doc.add_heading(match.group(1), level=level)
                # Check for table markers
                elif "<TABLE_MARKER>" in line and "</TABLE_MARKER>" in line:
                    table_html = re.search(r"<TABLE_MARKER>(.*?)</TABLE_MARKER>", line, re.DOTALL)
                    if table_html:
                        try:
                            html_table_to_docx(table_html.group(1), doc)
                        except Exception as e:
                            print(f"[J-MESSAGES PRE-EXPORT] ⚠️ Failed to convert table: {e}")
                            doc.add_paragraph("[Table conversion failed]")
                # Check for paragraphs
                elif line.startswith("<p"):
                    match = re.search(r">([^<]+)<", line)
                    if match:
                        doc.add_paragraph(match.group(1))
                else:
                    # Plain text paragraph
                    # Remove HTML tags
                    clean_text = re.sub(r"<[^>]+>", "", line)
                    if clean_text.strip():
                        doc.add_paragraph(clean_text.strip())
        else:
            # Fallback to raw_text
            for line in (body_text or "").split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
        
        # Stream as response
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        filename = f"j-message-pre-analyzed-{(data.get('id') or 'document')}.docx"
        return StreamingResponse(
            bio,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {e}")

